import datetime
from docx.api import Document
import libs.OpgSchedule
from libs.OpgSchedule import OPGLesson, OPGWeekDay, OpgSchedule
import json

file = open("ressources/name_map.json", "r")
name_map = json.loads(file.read())
file.close()

class ExcuseWriter:
    __author_sig = "....................................................................."
    __start_date_sig = "............................."
    __end_date_sig = "......................"
    __reason_sig = ".................................................................................................................."
    __outputfile:str
    __start_date:str
    __end_date:str
    __author:str
    __reason:str
    __excuse:libs.OpgSchedule.ExcusePeriod

    __doc:Document

    def __init__(self, author:str, excuse:libs.OpgSchedule.ExcusePeriod, input_file:str, output_file:str,reason:str=None, sign_date:datetime.date=datetime.datetime.now().date()):
        self.__author = author
        self.__start_date = excuse.getStartDate().strftime("%d.%m.%Y")
        self.__end_date = excuse.getEndDate().strftime("%d.%m.%Y")
        self.__reason = reason
        self.__excuse = excuse

        self.__doc = Document(input_file)
        self.__output_file = output_file
        self.__initSignDate(sign_date)
        self.__initDefaults()
        self.__initSchedule()

    def getDocument(self) -> Document:
        return self.__doc

    def safe(self):
        self.__doc.save(self.__output_file)

    def __initSchedule(self):
        schedule = self.__doc.tables[1]

        for abstence in self.__excuse.getWeekDays():
            x = abstence.day_count + 1
            for lesson in abstence.getOpgLessons():
                y = lesson.number + 1
                if y >= len(OPGWeekDay.lessons):
                    break
                cell = schedule.cell(y, x)
                if lesson.wasStudentPartOf():
                    cell.text = ""
                    continue
                cell.text = "\n" + name_map[lesson.getLessonName()]

    def __initDefaults(self):
        objs = []
        for paragraph in self.__doc.paragraphs:
            for run in paragraph.runs:
                objs.append(run)
        author_obj = objs[1]
        from_to_obj = objs[3]
        reason_obj = objs[4]

        author_obj.text = str(author_obj.text).replace(self.__author_sig, self.__author + self.__author_sig[len(self.__author) * 2:len(self.__author_sig)])
        from_to_obj.text = str(from_to_obj.text).replace(self.__start_date_sig, self.__start_date + self.__start_date_sig[len(self.__start_date) * 2 - 2:len(self.__start_date_sig)], 1)
        from_to_obj.text = str(from_to_obj.text).replace(self.__end_date_sig, self.__end_date + self.__end_date_sig[len(self.__end_date) * 2 - 2:len(self.__end_date_sig)])
        if self.__reason is not None:
            reason_obj.text = str(reason_obj.text).replace(self.__reason_sig, self.__reason + self.__reason_sig[len(self.__reason) * 2 - 1:len(self.__reason_sig)])

    def __initSignDate(self, date:datetime.date):
        self.__doc.tables[0].cell(0, 0).text = str(date.strftime("%d.%m.%Y"))
